#!/usr/bin/env python3
"""
文档API搜索器 - 在非C/C++文件中搜索API使用说明
"""

import os
import re
import logging
from typing import List
from .file_extensions import is_document_file, is_text_based_document
from .file_finder import FileFinder

# 条件导入PDF和Word文档解析库
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

logger = logging.getLogger(__name__)


class ApiDocumentInfo:
    """API文档信息类"""
    
    def __init__(self, api_name: str, file_path: str, line_number: int, 
                 context: str, match_type: str = "exact"):
        self.api_name = api_name
        self.file_path = file_path
        self.line_number = line_number
        self.context = context  # 包含API的上下文内容
        self.match_type = match_type  # exact, partial, description
        self.file_type = self._get_file_type()
        
    def _get_file_type(self) -> str:
        """获取文件类型"""
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext == '.md':
            return 'markdown'
        elif ext == '.rst':
            return 'restructuredtext'
        elif ext == '.txt':
            return 'text'
        elif ext in ['.doc', '.docx']:
            return 'word'
        elif ext == '.pdf':
            return 'pdf'
        else:
            return 'unknown'
    
    def to_dict(self) -> dict:
        """转换为字典"""
        return {
            'api_name': self.api_name,
            'file_path': self.file_path,
            'file_name': os.path.basename(self.file_path),
            'line_number': self.line_number,
            'context': self.context,
            'match_type': self.match_type,
            'file_type': self.file_type
        }


class DocumentApiSearcher:
    """文档API搜索器"""
    
    def __init__(self):
        self.file_finder = FileFinder()
        self.context_lines = 3  # 上下文行数
        
    def search_api_in_documents(self, api_name: str, search_path: str, 
                               recursive: bool = True, 
                               use_paragraph_extraction: bool = True,
                               target_files: List[str] = None) -> List[ApiDocumentInfo]:
        """
        在文档文件中搜索API使用说明
        
        Args:
            api_name: 要搜索的API名称
            search_path: 搜索路径
            recursive: 是否递归搜索
            use_paragraph_extraction: 是否使用段落提取，None时使用实例默认值
            target_files: 指定要搜索的文件列表，如果为None则搜索所有文档文件
            
        Returns:
            包含API信息的列表
        """
        results = []
        
        # 使用指定的提取方式
        
        # 查找文档文件
        if target_files is not None:
            # 使用指定的文件列表，将相对路径转换为绝对路径并过滤出文档文件
            doc_files = []
            for f in target_files:
                # 如果是相对路径，则相对于search_path解析
                if not os.path.isabs(f):
                    full_path = os.path.join(search_path, f)
                else:
                    full_path = f
                
                # 检查文件是否存在且是文档文件
                if os.path.exists(full_path) and is_document_file(full_path):
                    doc_files.append(full_path)
                else:
                    logger.warning(f"指定的文档文件不存在或不是文档文件: {full_path}")
        else:
            # 查找所有文档文件
            doc_files = self._find_document_files(search_path, recursive)
        
        logger.info(f"在 {len(doc_files)} 个文档文件中搜索API '{api_name}'")
        
        for file_path in doc_files:
            try:
                file_results = self._search_api_in_file(api_name, file_path, use_paragraph_extraction)
                results.extend(file_results)
            except Exception as e:
                logger.warning(f"搜索文件 {file_path} 时出错: {e}")
        
        logger.info(f"找到 {len(results)} 个API使用说明")
        return results
    
    def _find_document_files(self, search_path: str, recursive: bool) -> List[str]:
        """查找文档文件"""
        all_files = []
        
        if os.path.isfile(search_path):
            if is_document_file(search_path):
                return [search_path]
            else:
                return []
        
        # 遍历目录
        for root, dirs, files in os.walk(search_path):
            for file in files:
                file_path = os.path.join(root, file)
                if is_document_file(file_path):
                    all_files.append(file_path)
            
            if not recursive:
                break
        
        return all_files
    
    def _search_api_in_file(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """在单个文件中搜索API"""
        results = []
        
        # 处理不同类型的文档文件
        if is_text_based_document(file_path):
            # 处理文本格式文档
            return self._search_text_document(api_name, file_path, use_paragraph_extraction)
        else:
            # 处理二进制格式文档（PDF、DOC等）
            return self._search_binary_document(api_name, file_path, use_paragraph_extraction)
    
    def _search_text_document(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """在文本格式文档中搜索API"""
        results = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
        except Exception as e:
            logger.warning(f"无法读取文件 {file_path}: {e}")
            return results
        
        # 搜索API名称
        for i, line in enumerate(lines):
            line_number = i + 1
            
            # 检查是否包含API名称
            if api_name.lower() in line.lower():
                # 优先检查精确匹配
                if self._is_exact_match(api_name, line):
                    match_type = "exact"
                else:
                    match_type = "partial"
                
                context = self._extract_context(lines, i, use_paragraph_extraction)
                results.append(ApiDocumentInfo(
                    api_name=api_name,
                    file_path=file_path,
                    line_number=line_number,
                    context=context,
                    match_type=match_type
                ))
        
        return results
    
    def _search_binary_document(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """在二进制格式文档中搜索API（PDF、DOC等）"""
        results = []
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._search_pdf_document(api_name, file_path, use_paragraph_extraction)
            elif file_ext in ['.doc', '.docx']:
                return self._search_word_document(api_name, file_path, use_paragraph_extraction)
            else:
                logger.debug(f"暂不支持的文档格式: {file_ext}")
                return results
        except Exception as e:
            logger.warning(f"处理二进制文档 {file_path} 时出错: {e}")
            return results
    
    def _search_pdf_document(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """在PDF文档中搜索API"""
        if HAS_PYPDF2:
            return self._search_pdf_with_pypdf2(api_name, file_path, use_paragraph_extraction)
        elif HAS_PDFPLUMBER:
            return self._search_pdf_with_pdfplumber(api_name, file_path, use_paragraph_extraction)
        else:
            logger.warning(f"未安装PDF解析库（PyPDF2或pdfplumber），跳过PDF文件: {file_path}")
            return []
    
    def _search_word_document(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """在Word文档中搜索API"""
        if HAS_PYTHON_DOCX:
            return self._search_docx_with_python_docx(api_name, file_path, use_paragraph_extraction)
        else:
            logger.warning(f"未安装python-docx库，跳过Word文件: {file_path}")
            return []
    
    def _search_pdf_with_pypdf2(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """使用PyPDF2解析PDF文档"""
        results = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        lines = text.split('\n')
                        
                        for line_num, line in enumerate(lines, 1):
                            if self._is_exact_match(api_name, line):
                                context = self._extract_context_from_text(text, line)
                                results.append(ApiDocumentInfo(
                                    api_name=api_name,
                                    file_path=file_path,
                                    line_number=f"Page {page_num}, Line {line_num}",
                                    context=context,
                                    match_type="exact"
                                ))
                            elif api_name.lower() in line.lower():
                                context = self._extract_context_from_text(text, line)
                                results.append(ApiDocumentInfo(
                                    api_name=api_name,
                                    file_path=file_path,
                                    line_number=f"Page {page_num}, Line {line_num}",
                                    context=context,
                                    match_type="partial"
                                ))
                    except Exception as e:
                        logger.warning(f"解析PDF第{page_num}页时出错: {e}")
                        continue
        except Exception as e:
            logger.warning(f"读取PDF文件 {file_path} 时出错: {e}")
        
        return results
    
    def _search_pdf_with_pdfplumber(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """使用pdfplumber解析PDF文档"""
        results = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        text = page.extract_text()
                        if text:
                            lines = text.split('\n')
                            
                            for line_num, line in enumerate(lines, 1):
                                if self._is_exact_match(api_name, line):
                                    context = self._extract_context_from_text(text, line)
                                    results.append(ApiDocumentInfo(
                                        api_name=api_name,
                                        file_path=file_path,
                                        line_number=f"Page {page_num}, Line {line_num}",
                                        context=context,
                                        match_type="exact"
                                    ))
                                elif api_name.lower() in line.lower():
                                    context = self._extract_context_from_text(text, line)
                                    results.append(ApiDocumentInfo(
                                        api_name=api_name,
                                        file_path=file_path,
                                        line_number=f"Page {page_num}, Line {line_num}",
                                        context=context,
                                        match_type="partial"
                                    ))
                    except Exception as e:
                        logger.warning(f"解析PDF第{page_num}页时出错: {e}")
                        continue
        except Exception as e:
            logger.warning(f"读取PDF文件 {file_path} 时出错: {e}")
        
        return results
    
    def _search_docx_with_python_docx(self, api_name: str, file_path: str, use_paragraph_extraction: bool = True) -> List[ApiDocumentInfo]:
        """使用python-docx解析Word文档"""
        results = []
        
        try:
            doc = Document(file_path)
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                text = paragraph.text
                if self._is_exact_match(api_name, text):
                    context = self._extract_context_from_paragraphs(doc.paragraphs, para_num - 1)
                    results.append(ApiDocumentInfo(
                        api_name=api_name,
                        file_path=file_path,
                        line_number=f"Paragraph {para_num}",
                        context=context,
                        match_type="exact"
                    ))
                elif api_name.lower() in text.lower():
                    context = self._extract_context_from_paragraphs(doc.paragraphs, para_num - 1)
                    results.append(ApiDocumentInfo(
                        api_name=api_name,
                        file_path=file_path,
                        line_number=f"Paragraph {para_num}",
                        context=context,
                        match_type="partial"
                    ))
        except Exception as e:
            logger.warning(f"读取Word文件 {file_path} 时出错: {e}")
        
        return results
    
    def _extract_context_from_text(self, text: str, target_line: str) -> str:
        """从文本中提取上下文"""
        lines = text.split('\n')
        target_index = -1
        
        for i, line in enumerate(lines):
            if line.strip() == target_line.strip():
                target_index = i
                break
        
        if target_index == -1:
            return target_line
        
        # 提取前后各2行作为上下文
        start = max(0, target_index - 2)
        end = min(len(lines), target_index + 3)
        context_lines = lines[start:end]
        
        return '\n'.join(context_lines)
    
    def _extract_context_from_paragraphs(self, paragraphs, target_index: int) -> str:
        """从段落列表中提取上下文"""
        # 提取前后各1个段落作为上下文
        start = max(0, target_index - 1)
        end = min(len(paragraphs), target_index + 2)
        
        context_paragraphs = []
        for i in range(start, end):
            if paragraphs[i].text.strip():
                context_paragraphs.append(paragraphs[i].text)
        
        return '\n\n'.join(context_paragraphs)
    
    def _is_exact_match(self, api_name: str, line: str) -> bool:
        """检查是否为精确匹配"""
        # 使用正则表达式进行单词边界匹配
        pattern = r'\b' + re.escape(api_name) + r'\b'
        return bool(re.search(pattern, line, re.IGNORECASE))
    
    def _extract_context(self, lines: List[str], target_line_index: int, use_paragraph_extraction: bool = True) -> str:
        """提取上下文"""
        if use_paragraph_extraction:
            # 尝试基于段落提取，如果失败则回退到固定行数
            paragraph_context = self._extract_paragraph_context(lines, target_line_index)
            if paragraph_context:
                return paragraph_context
            
            # 回退到固定行数提取
            return self._extract_fixed_lines_context(lines, target_line_index)
        else:
            # 直接使用固定行数提取
            return self._extract_fixed_lines_context(lines, target_line_index)
    
    def _extract_paragraph_context(self, lines: List[str], target_line_index: int) -> str:
        """基于段落边界提取上下文"""
        # 向上查找段落开始
        start = target_line_index
        while start > 0 and lines[start - 1].strip() != "":
            start -= 1
        
        # 向下查找段落结束
        end = target_line_index
        while end < len(lines) - 1 and lines[end + 1].strip() != "":
            end += 1
        
        # 如果段落太长（超过10行），限制范围
        if end - start > 10:
            # 以目标行为中心，取前后各5行
            start = max(0, target_line_index - 5)
            end = min(len(lines) - 1, target_line_index + 5)
        
        # 如果段落太短（只有1行），扩展到包含前后段落
        if end == start:
            # 向上扩展一个段落
            while start > 0 and lines[start - 1].strip() == "":
                start -= 1
            while start > 0 and lines[start - 1].strip() != "":
                start -= 1
            
            # 向下扩展一个段落
            while end < len(lines) - 1 and lines[end + 1].strip() == "":
                end += 1
            while end < len(lines) - 1 and lines[end + 1].strip() != "":
                end += 1
        
        context_lines = []
        for i in range(start, end + 1):
            line_num = i + 1
            line_content = lines[i].rstrip()
            if i == target_line_index:
                # 标记目标行
                context_lines.append(f">>> {line_num:4d}: {line_content}")
            else:
                context_lines.append(f"    {line_num:4d}: {line_content}")
        
        return '\n'.join(context_lines)
    
    def _extract_fixed_lines_context(self, lines: List[str], target_line_index: int) -> str:
        """固定行数提取上下文（回退方案）"""
        start = max(0, target_line_index - self.context_lines)
        end = min(len(lines), target_line_index + self.context_lines + 1)
        
        context_lines = []
        for i in range(start, end):
            line_num = i + 1
            line_content = lines[i].rstrip()
            if i == target_line_index:
                # 标记目标行
                context_lines.append(f">>> {line_num:4d}: {line_content}")
            else:
                context_lines.append(f"    {line_num:4d}: {line_content}")
        
        return '\n'.join(context_lines)